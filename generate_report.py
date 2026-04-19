import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement, ns

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(ns.qn(name), value)

def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'separate')

    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

doc = Document()

# Set to A4 (8.27 x 11.69 inches)
section = doc.sections[0]
section.page_width = Inches(8.27)
section.page_height = Inches(11.69)

# Add page number in footer (Bottom-Right)
footer = section.footer
footer_para = footer.paragraphs[0]
footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_page_number(footer_para.add_run())

# Define Styles
styles = doc.styles

# Heading Style (15pt Arial)
heading_style = styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
heading_style.font.name = 'Arial'
heading_style.font.size = Pt(15)
heading_style.font.bold = True

# Subheading Style (14pt Arial)
subheading_style = styles.add_style('CustomSubheading', WD_STYLE_TYPE.PARAGRAPH)
subheading_style.font.name = 'Arial'
subheading_style.font.size = Pt(14)
subheading_style.font.bold = True

# Body Style (12pt Arial)
body_style = styles.add_style('CustomBody', WD_STYLE_TYPE.PARAGRAPH)
body_style.font.name = 'Arial'
body_style.font.size = Pt(12)
body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_heading(text):
    p = doc.add_paragraph(text, style='CustomHeading')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_subheading(text):
    doc.add_paragraph(text, style='CustomSubheading')

def add_body(text):
    p = doc.add_paragraph(text, style='CustomBody')
    return p

# 1. Title
add_heading("InnovaCorp O2C: Advanced Order-to-Cash Automation Platform")
doc.add_paragraph() # Spacing

# 2. Student Details
details = doc.add_paragraph(style='CustomBody')
details.add_run("Name: ").bold = True
details.add_run("Saket Suman\n")
details.add_run("Roll Number: ").bold = True
details.add_run("23052419\n")
details.add_run("Batch/Programme: ").bold = True
details.add_run("SAP BTP Developer")
doc.add_paragraph()

# 3. Problem Statement
add_subheading("Problem Statement")
add_body("In modern enterprise sales cycles, the Order-to-Cash (O2C) process is often hindered by manual order approvals, disjointed credit block resolutions, and inefficient dunning processes. Traditional systems lack real-time visibility and automated workflows, resulting in delayed cash flows, increased financial risk, and a suboptimal customer experience. A unified, automated solution is required to seamlessly handle credit checks, streamline manager approvals, and manage outstanding payments to ensure uninterrupted business operations.")
doc.add_paragraph()

# 4. Solution/Features
add_subheading("Solution/Features")
add_body("InnovaCorp O2C is a comprehensive Order-to-Cash application built using the SAP Cloud Application Programming (CAP) model and deployed on the SAP Business Technology Platform (BTP). It automates and optimizes the entire sales lifecycle from order creation to payment collection. Key features include:")
add_body("• Automated Sales Order Processing: Real-time credit evaluation upon order submission.\n"
         "• Dynamic Credit Block Logging: Intelligent interception of orders exceeding customer credit limits, securely logging excess amounts and restricting further processing.\n"
         "• SAP Build Process Automation (SBPA) Integration: Non-blocking, event-driven workflow orchestration for multi-level credit manager approvals via a centralized inbox.\n"
         "• Automated Dunning & Notifications: Tiered automated dunning levels (Reminder, Warning, Final) based on payment overdue days, leveraging robust notification services.")
doc.add_paragraph()

# 5. Screenshots
add_subheading("Screenshots")
add_body("[ Please insert Screenshot 1 here: Customer Credit Overview Dashboard ]")
add_body("[ Please insert Screenshot 2 here: Sales Orders List / Statuses ]")
add_body("[ Please insert Screenshot 3 here: InnovaCorp Solutions Generated INVOICE ]")
add_body("[ Please insert Screenshot 4 here: New Sales Order Modal (Credit Limit Warning) ]")
add_body("[ Please insert Screenshot 5 here: O2C Analytics KPI Dashboard ]")
doc.add_paragraph()

# 6. Tech Stack
add_subheading("Tech Stack")
add_body("• Backend & Framework: SAP CAP (Cloud Application Programming Model), Node.js, Express\n"
         "• Database: SAP HANA Cloud / local SQLite for development data modeling\n"
         "• Workflow & Orchestration: SAP Build Process Automation (SBPA)\n"
         "• Cloud Platform: SAP BTP (Business Technology Platform) Trial Environment\n"
         "• API / Integration: REST APIs, @sap/cds modules")
doc.add_paragraph()

# 7. Unique Points
add_subheading("Unique Points")
add_body("• Event-Driven Architecture: Deep integration with SBPA using non-blocking API triggers ensuring high system responsiveness without halting the main transaction thread.\n"
         "• Comprehensive Concept-to-Code: The application efficiently translates the conceptual order-to-cash business pipeline into clean executable code using the advanced SAP CDS structure.\n"
         "• Extensible Modular Services: The design encapsulates the notification pipeline, providing a clean separation of concerns and simpler scalability for future enterprise notifications.")
doc.add_paragraph()

# 8. Future Improvements
add_subheading("Future Improvements")
add_body("• AI-Driven Risk Assessment: Integration with Machine Learning models deployed on SAP AI Core to calculate dynamic customer credit scores and default probabilities.\n"
         "• Multi-Currency Payment Integration: Real-time FX conversion handling combined with major ERP integrations (like S/4HANA) to close the loop on cash collection.\n"
         "• Advanced Analytics & Dashboard: Coupling with SAP Analytics Cloud (SAC) to present real-time telemetry on cash flow constraints and process bottlenecks.")

# Save the document
doc.save("Project_Report_InnovaCorp_O2C_v2.docx")
